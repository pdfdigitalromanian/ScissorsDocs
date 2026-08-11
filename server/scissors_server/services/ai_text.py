from __future__ import annotations

import re
from collections import Counter
from pathlib import Path
from typing import Any, Dict, List, Mapping, Sequence

from ..models import ToolInputError, ToolResult
from .common import option_int, option_text, require_file, require_files

WORD_PATTERN = re.compile(r"[A-Za-zÀ-ÖØ-öø-ÿ0-9']+")
SENTENCE_PATTERN = re.compile(r"(?<=[.!?])\s+")
STOP_WORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "for",
    "from",
    "has",
    "in",
    "is",
    "it",
    "of",
    "on",
    "or",
    "that",
    "the",
    "this",
    "to",
    "was",
    "were",
    "will",
    "with",
}


def _extract_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        import pymupdf

        document = pymupdf.open(str(path))
        text = "\n\n".join(page.get_text("text") for page in document)
        document.close()
        return text
    if suffix == ".docx":
        from docx import Document

        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            parts.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(parts)
    if suffix in {".html", ".htm"}:
        from bs4 import BeautifulSoup

        return BeautifulSoup(
            path.read_text(encoding="utf-8", errors="replace"),
            "html.parser",
        ).get_text("\n", strip=True)
    return path.read_text(encoding="utf-8", errors="replace")


def _request_text(files: Sequence[Path], options: Mapping[str, Any]) -> str:
    if files:
        text = _extract_text(files[0]).strip()
    else:
        text = option_text(options, "text", required=True)
    if not text:
        raise ToolInputError("The document contains no extractable text.")
    return text


def summarize(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    del workdir
    text = _request_text(files, options)
    sentences = [sentence.strip() for sentence in SENTENCE_PATTERN.split(text) if sentence.strip()]
    sentence_count = option_int(options, "sentences", 5, minimum=1, maximum=20)
    if len(sentences) <= sentence_count:
        summary = " ".join(sentences)
    else:
        words = [word.lower() for word in WORD_PATTERN.findall(text)]
        frequencies = Counter(word for word in words if word not in STOP_WORDS and len(word) > 2)
        ranked = sorted(
            enumerate(sentences),
            key=lambda item: sum(
                frequencies[word.lower()] for word in WORD_PATTERN.findall(item[1])
            ),
            reverse=True,
        )[:sentence_count]
        summary = " ".join(sentence for _, sentence in sorted(ranked))
    return ToolResult.json(
        {
            "summary": summary,
            "source_characters": len(text),
            "summary_characters": len(summary),
            "method": "local-extractive",
        }
    )


def translate(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    del workdir
    from deep_translator import GoogleTranslator

    text = _request_text(files, options)
    source = option_text(options, "source", "auto")
    target = option_text(options, "target", required=True)
    chunks = [text[index : index + 4500] for index in range(0, len(text), 4500)]
    try:
        translated = "".join(
            GoogleTranslator(source=source, target=target).translate(chunk) or ""
            for chunk in chunks
        )
    except Exception as error:
        raise ToolInputError(
            "Translation service could not be reached. Check the language codes and network."
        ) from error
    return ToolResult.json(
        {
            "translated_text": translated,
            "source_language": source,
            "target_language": target,
            "method": "deep-translator",
        }
    )


def rewrite(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    del workdir
    text = _request_text(files, options)
    mode = option_text(options, "mode", "concise").lower()
    replacements: Dict[str, str]
    if mode == "formal":
        replacements = {
            r"\bcan't\b": "cannot",
            r"\bwon't\b": "will not",
            r"\bdon't\b": "do not",
            r"\bget\b": "obtain",
            r"\ba lot of\b": "many",
        }
    elif mode == "plain":
        replacements = {
            r"\butilize\b": "use",
            r"\bapproximately\b": "about",
            r"\bcommence\b": "start",
            r"\bterminate\b": "end",
            r"\bin order to\b": "to",
        }
    else:
        replacements = {
            r"\bin order to\b": "to",
            r"\bdue to the fact that\b": "because",
            r"\bat this point in time\b": "now",
            r"\bfor the purpose of\b": "for",
            r"\bvery\s+": "",
        }
    rewritten = text
    for pattern, replacement in replacements.items():
        rewritten = re.sub(
            pattern,
            lambda match, value=replacement: (
                value.capitalize() if match.group(0)[0].isupper() else value
            ),
            rewritten,
            flags=re.IGNORECASE,
        )
    rewritten = re.sub(r"[ \t]+", " ", rewritten)
    return ToolResult.json(
        {
            "rewritten_text": rewritten.strip(),
            "mode": mode,
            "method": "local-rules",
        }
    )


def assistant(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    del workdir
    text = _request_text(files, options)
    question = option_text(options, "question", required=True)
    keywords = {
        word.lower()
        for word in WORD_PATTERN.findall(question)
        if word.lower() not in STOP_WORDS and len(word) > 2
    }
    sentences = [sentence.strip() for sentence in SENTENCE_PATTERN.split(text) if sentence.strip()]
    matches = sorted(
        sentences,
        key=lambda sentence: len(
            keywords.intersection(word.lower() for word in WORD_PATTERN.findall(sentence))
        ),
        reverse=True,
    )
    relevant = [
        sentence
        for sentence in matches[:3]
        if keywords.intersection(word.lower() for word in WORD_PATTERN.findall(sentence))
    ]
    answer = (
        " ".join(relevant)
        if relevant
        else "No directly relevant passage was found in the document."
    )
    return ToolResult.json(
        {
            "answer": answer,
            "question": question,
            "method": "local-retrieval",
        }
    )


def extract_tables(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    del workdir
    import pymupdf

    pdf = require_file(require_files(files), {".pdf"})
    document = pymupdf.open(str(pdf))
    tables: List[Dict[str, Any]] = []
    for page_number, page in enumerate(document, start=1):
        found = page.find_tables()
        for table_number, table in enumerate(found.tables, start=1):
            tables.append(
                {
                    "page": page_number,
                    "table": table_number,
                    "rows": table.extract(),
                }
            )
    document.close()
    return ToolResult.json(
        {
            "tables": tables,
            "count": len(tables),
            "method": "pymupdf-table-detection",
        }
    )
