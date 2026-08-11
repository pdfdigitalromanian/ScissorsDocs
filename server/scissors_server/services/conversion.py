from __future__ import annotations

from pathlib import Path
from typing import Any, List, Mapping, Sequence
from xml.sax.saxutils import escape

from ..models import ToolInputError, ToolResult
from .common import (
    DOCX_MEDIA_TYPE,
    PDF_MEDIA_TYPE,
    ZIP_MEDIA_TYPE,
    create_zip,
    option_int,
    option_text,
    output_path,
    require_file,
    require_files,
    require_pdf,
)


def images_to_pdf(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    from PIL import Image, ImageOps

    image_paths = [
        path
        for path in require_files(files)
        if path.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}
    ]
    if not image_paths:
        raise ToolInputError("Add at least one PNG, JPEG, WEBP, or TIFF image.")
    images: List[Any] = []
    try:
        for path in image_paths:
            with Image.open(path) as source:
                images.append(ImageOps.exif_transpose(source).convert("RGB"))
        destination = output_path(workdir, "images", ".pdf")
        images[0].save(
            destination,
            "PDF",
            resolution=option_int(options, "dpi", 150, minimum=72, maximum=600),
            save_all=True,
            append_images=images[1:],
        )
    finally:
        for image in images:
            image.close()
    return ToolResult.file(destination, PDF_MEDIA_TYPE)


def pdf_to_images(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    import pymupdf

    pdf = require_pdf(files)
    document = pymupdf.open(str(pdf))
    dpi = option_int(options, "dpi", 150, minimum=72, maximum=600)
    image_format = option_text(options, "format", "png").lower()
    if image_format not in {"png", "jpg", "jpeg"}:
        document.close()
        raise ToolInputError("Image format must be PNG or JPEG.")
    suffix = "jpg" if image_format in {"jpg", "jpeg"} else "png"
    outputs = []
    for index, page in enumerate(document, start=1):
        destination = output_path(workdir, f"{pdf.stem}-page-{index}", f".{suffix}")
        pixmap = page.get_pixmap(dpi=dpi, alpha=False)
        pixmap.save(str(destination))
        outputs.append(destination)
    document.close()
    archive = create_zip(outputs, output_path(workdir, f"{pdf.stem}-images", ".zip"))
    return ToolResult.file(archive, ZIP_MEDIA_TYPE)


def word_to_pdf(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    from docx import Document
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    source = require_file(files, {".docx"})
    document = Document(str(source))
    destination = output_path(workdir, f"{source.stem}-converted", ".pdf")
    page_size = LETTER if option_text(options, "page_size", "a4").lower() == "letter" else A4
    pdf = SimpleDocTemplate(str(destination), pagesize=page_size)
    styles = getSampleStyleSheet()
    story = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            story.append(Paragraph(escape(paragraph.text), styles["BodyText"]))
            story.append(Spacer(1, 8))
    for table in document.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if rows:
            rendered = Table(rows, repeatRows=1)
            rendered.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.append(rendered)
            story.append(Spacer(1, 12))
    if not story:
        story.append(Paragraph(" ", styles["BodyText"]))
    pdf.build(story)
    return ToolResult.file(destination, PDF_MEDIA_TYPE)


def pdf_to_word(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    import pymupdf
    from docx import Document

    pdf = require_pdf(files)
    source = pymupdf.open(str(pdf))
    document = Document()
    document.core_properties.title = pdf.stem
    for index, page in enumerate(source):
        if index:
            document.add_page_break()
        text = page.get_text("text").strip()
        if not text:
            document.add_paragraph("[Page contains no extractable text]")
            continue
        for block in text.split("\n\n"):
            document.add_paragraph(block.strip())
    source.close()
    destination = output_path(workdir, f"{pdf.stem}-converted", ".docx")
    document.save(str(destination))
    return ToolResult.file(destination, DOCX_MEDIA_TYPE)


def html_to_pdf(
    files: Sequence[Path], options: Mapping[str, Any], workdir: Path
) -> ToolResult:
    from bs4 import BeautifulSoup
    from reportlab.lib.pagesizes import A4, LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    html_path = None
    for path in files:
        if path.suffix.lower() in {".html", ".htm"}:
            html_path = path
            break
    if html_path:
        html = html_path.read_text(encoding="utf-8", errors="replace")
        stem = html_path.stem
    else:
        html = option_text(options, "html", required=True)
        stem = "web-page"
    soup = BeautifulSoup(html, "html.parser")
    for element in soup(["script", "style", "noscript"]):
        element.decompose()
    title = soup.title.string.strip() if soup.title and soup.title.string else stem
    destination = output_path(workdir, f"{stem}-converted", ".pdf")
    page_size = LETTER if option_text(options, "page_size", "a4").lower() == "letter" else A4
    pdf = SimpleDocTemplate(str(destination), pagesize=page_size, title=title)
    styles = getSampleStyleSheet()
    story = [Paragraph(escape(title), styles["Title"]), Spacer(1, 12)]
    for line in soup.get_text("\n", strip=True).splitlines():
        if line.strip() and line.strip() != title:
            story.append(Paragraph(escape(line.strip()), styles["BodyText"]))
            story.append(Spacer(1, 6))
    pdf.build(story)
    return ToolResult.file(destination, PDF_MEDIA_TYPE)
